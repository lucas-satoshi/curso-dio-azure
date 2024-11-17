import requests
from docx import Document
import os
from bs4 import BeautifulSoup
from langchain_openai.chat_models.azure import AzureChatOpenAI

subscription_key = "006a49775d0e4eca80c84d7277ccd653"
endpoint = 'https://api.cognitive.microsofttranslator.com'
location = "eastus2"
language_destination = 'pt-br'

client = AzureChatOpenAI(
    azure_endpoint= "https://oai-dio-bootcamp-dev-eastus-001.openai.azure.com/",
    api_key= "730f2276a6a54c01964cf11672f08e7c",
    api_version= "2024-02-15-preview",
    deployment_name= "gpt-4o-mini",
    max_retries=0
)

def translator_text(text, target_language):
  path = '/translate'
  constructed_url = endpoint + path
  headers = {
      'Ocp-Apim-Subscription-Key': subscription_key,
      'Ocp-Apim-Subscription-Region': location,
      'Content-type': 'application/json',
      'X-ClientTraceId': str(os.urandom(16))
  }

  body = [{
      'text' : text
  }]
  params = {
      'api-version' : '3.0',
      'from' : 'en',
      'to' : target_language
  }
  request = requests.post(constructed_url, params=params, headers=headers, json=body)
  response = request.json()

  # Check if the response is a list and contains the expected structure
  if isinstance(response, list) and len(response) > 0 and "translations" in response[0] and len(response[0]["translations"]) > 0 and "text" in response[0]["translations"][0]:
    return response[0]["translations"][0]["text"]
  else:
    # Handle the case where the expected structure is not present
    print("Error: Unexpected response format from the API.")
    print("Response:", response)  # Print the response for debugging
    return None  # Or raise an exception, depending on your needs
  
def translate_document(path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    translated_text = translator_text(paragraph.text, language_destination)
    full_text.append(translated_text)

  translated_doc = Document()
  for line in full_text:
    translated_doc.add_paragraph(line)
  path_translated = path.replace("docx", f"{language_destination}.docx")
  translated_doc.save(path_translated)
  return path_translated

def extract_text_from_url(url):
  response = requests.get(url)

  if response.status_code != 200:
    soup = BeautifulSoup(response.text, 'html.parser')
    for script_or_style in soup(["script", "style"]):
      script_or_style.decompose()
    texto = soup.get_text(separator= ' ')
    #Limpar text
    linhas = (line.strip() for line in texto.splitlines())
    parts = (phrase.strip() for line in linhas for phrase in line.split("  "))
    texto_limpo = '\n'.join(part for part in parts if part)
    return texto_limpo
  else:
    print(f"Failed to fetch the URL. Status code: {response.status_code}")
    return None

  
  text = soup.get_text()
  return text

def translate_article(text, lang):
  messages = [
      ("system", "Você atua como tradutor de textos"),
      ("user", f"Traduza o {text} para o idioma {lang} e responda em markdown")
  ]

  response = client.invoke(messages)
  print(response.content)
  return response.content

translator_text('I know you are somewhere out there, somewhere far away', language_destination)

input_file = "/content/MUSICA.docx"
translate_document(input_file)

extract_text_from_url('https://dev.to/kenakamu/azure-open-ai-in-vnet-3alo')

translate_article("Let's see if the deployment was succeeded.", "portugues")

url = 'https://dev.to/kenakamu/azure-open-ai-in-vnet-3alo'
text = extract_text_from_url(url)
article = translate_article(text, "pt-br")

print(article)